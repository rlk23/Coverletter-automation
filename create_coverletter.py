import os
from docx import Document
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.schema import Document as LangchainDocument
from bs4 import BeautifulSoup
import requests
import re
from splitter import text_to_doc_splitter, split_text_documents
from fileReader import load_document
from web_scraper import extract_text_from_url

# Function to extract current position from resume text
def extract_current_position(resume_text):
    match = re.search(r'(\b[A-Za-z ]+\b) at \b[A-Za-z ]+\b', resume_text)
    if match:
        return match.group(1)
    return "Current Position"

# Function to extract job title from job description
def extract_job_title(job_description):
    match = re.search(r'\b(Job Title|Position):\s*(\b[A-Za-z ]+\b)', job_description)
    if match:
        return match.group(2)
    return "Job Title"

# Function to create the cover letter based on prompts and extracted information
def create_cover_letter(resume_text, job_description, current_position, job_title, openai_api_key):
    pdf_doc = [LangchainDocument(page_content=resume_text)]
    pdf_doc.append(LangchainDocument(page_content=job_description))
    documents = split_text_documents(pdf_doc)

    vectordb = Chroma.from_documents(documents, embedding=OpenAIEmbeddings(openai_api_key=openai_api_key))

    pdf_qa = RetrievalQA.from_chain_type(
        ChatOpenAI(temperature=0.7, model_name='gpt-3.5-turbo', openai_api_key=openai_api_key),
        retriever=vectordb.as_retriever(search_kwargs={'k': 6}),
        chain_type="stuff",
    )

    prompts = [
        "Identify the biggest day-to-day challenge someone in this position would face based on the provided job description.",
        f"You are currently a {current_position} and you are applying for a {job_title}. Write an attention grabbing hook for your cover letter that highlights your experience and qualifications, emphasizing your ability to empathize and tackle the challenges of the {job_title} role. Incorporate specific examples of how you've addressed similar challenges in your past work. Limit your response to 100 words.",
        "Complete the cover letter using the provided hook and resume. Ensure the cover letter is comprehensive yet concise, within 250 words."
    ]

    cover_letter_parts = []
    for prompt in prompts:
        result = pdf_qa.run(prompt)
        cover_letter_parts.append(result)

    cover_letter = "\n\n".join(cover_letter_parts)

    return cover_letter

# Function to create a Word document for the cover letter
def create_cover_letter_doc(cover_letter_text, output_path):
    document = Document()
    document.add_heading("Cover Letter", 0)
    document.add_paragraph(cover_letter_text)
    document.save(output_path)
    print(f"Cover letter saved to {output_path}")

# Main function to orchestrate the cover letter generation process
def generate_cover_letter(job_link, resume_path, openai_api_key):
    # Load resume document and extract text
    resume_doc = load_document(resume_path)
    resume_text = "\n".join([chunk.page_content for chunk in resume_doc])

    # Scrape job description from URL and extract text
    job_description_doc = extract_text_from_url(job_link)
    job_description = "\n".join([chunk.page_content for chunk in job_description_doc])

    # Extract current position from resume
    current_position = extract_current_position(resume_text)

    # Extract job title from job description
    job_title = extract_job_title(job_description)

    # Generate cover letter
    cover_letter_text = create_cover_letter(resume_text, job_description, current_position, job_title, openai_api_key)
    
    # Save cover letter to a Word document
    output_path = 'cover_letter.docx'
    create_cover_letter_doc(cover_letter_text, output_path)

    return output_path




