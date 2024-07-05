import fitz  # PyMuPDF
from PyPDF2 import PdfReader
from docx import Document
import os
from bs4 import BeautifulSoup
import requests
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.schema import Document as LangchainDocument
import re
from splitter import text_to_doc_splitter, split_text_documents

def extract_text(file_path):
    _, file_extension = os.path.splitext(file_path)

    if file_extension.lower() == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension.lower() == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError('Unsupported file type: {}'.format(file_extension))

def extract_text_from_pdf(file_path):
    pdf_reader = PdfReader(file_path)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    document = Document(file_path)
    text = "\n".join([para.text for para in document.paragraphs])
    return text

def extract_text_from_url(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        text = '\n'.join(
            section.get_text(separator='\n', strip=True)
            for section in soup.findAll('div', {'class': 'description__text'})
        )
        return text
    else:
        print(f"Failed to retrieve the page. Status code: {response.status_code}")
        return None

def extract_current_position(pdf_text):
    match = re.search(r'(\b[A-Za-z ]+\b) at \b[A-Za-z ]+\b', pdf_text)
    if match:
        return match.group(1)
    return "your current position"

def extract_job_title(job_posting_text):
    match = re.search(r'\b(Job Title|Position):\s*(\b[A-Za-z ]+\b)', job_posting_text)
    if match:
        return match.group(2)
    return "the position you are applying for"

def get_cover_letter(url, pdf, openai_api_key):
    pdf_text = extract_text(pdf)
    job_posting_text = extract_text_from_url(url)

    if job_posting_text is None:
        return "Failed to retrieve job posting text."

    current_position = extract_current_position(pdf_text)
    job_title = extract_job_title(job_posting_text)

    pdf_doc = [LangchainDocument(page_content=pdf_text)]
    pdf_doc.append(LangchainDocument(page_content=job_posting_text))
    documents = split_text_documents(pdf_doc)

    vectordb = Chroma.from_documents(documents, embedding=OpenAIEmbeddings(openai_api_key=openai_api_key))

    pdf_qa = RetrievalQA.from_chain_type(
        ChatOpenAI(temperature=0.7, model_name='gpt-3.5-turbo', openai_api_key=openai_api_key),
        retriever=vectordb.as_retriever(search_kwargs={'k': 6}),
        chain_type="stuff",
    )

    prompts = [
        "Identify the biggest day-to-day challenge someone in this position would face based on the provided job description.",
        f"You are currently a {current_position} and you are applying for a {job_title}. Write an attention grabbing hook for your cover letter that highlights your experience and qualifications, emphasizing your ability to empathize and tackle the challenges of the {job_title} role. Incorporate specific examples of how you've addressed similar challenges in your past work. Limit your response to 100 words.",
        "Complete the cover letter using the provided hook and resume. Ensure the cover letter is comprehensive yet concise, within 250 words."
    ]

    cover_letter_parts = []
    for prompt in prompts:
        result = pdf_qa.run(prompt)
        cover_letter_parts.append(result)

    cover_letter = "\n\n".join(cover_letter_parts)

    return cover_letter

if __name__ == "__main__":
    url = 'https://www.linkedin.com/jobs/view/3961701899/?alternateChannel=search&refId=1mEYxhdBPlDvxEKQCuDTHQ%3D%3D&trackingId=63WyeKH5AgkhI8h0EAqplg%3D%3D'
    pdf_path = ""
    openai_api_key = ""

    cover_letter = get_cover_letter(url, pdf_path, openai_api_key)
    print("Generated Cover Letter:\n", cover_letter)



